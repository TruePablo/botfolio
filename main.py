import telebot
import docx
import time
import schedule
global admin_id
global doc
doc = docx.Document()
from multiprocessing.context import Process

class ScheduleMessage():
    def try_send_schedule(s):
        while True:
            schedule.run_pending()
            time.sleep(1)

    def start_process(s):
        p1 = Process(target=ScheduleMessage.try_send_schedule, args=())
        p1.start()

TOKEN = '5198013721:AAGa238Qc8ckj3E2WZOwJyEHLdveQCfseDo'
bot = telebot.TeleBot(TOKEN)
admin_id = 615467250


keyboardanswer = telebot.types.ReplyKeyboardMarkup(True, True)
keyboardanswer.row('Да', 'Нет')

keyboardstrike = telebot.types.ReplyKeyboardMarkup(True, True)
keyboardstrike.row('Выгрузить страйки', 'Назад')

keyboardstartadmin = telebot.types.ReplyKeyboardMarkup(True, True)
keyboardstartadmin.row('Рассылка', 'Резюме', 'Страйки')

keyboardstart = telebot.types.ReplyKeyboardMarkup(True, True)
keyboardstart.row('Мои страйки', 'Резюме')

keyboardboss = telebot.types.ReplyKeyboardMarkup(True, True)
keyboardboss.row('Отправить резюме', 'Назад')

keyboardsend = telebot.types.ReplyKeyboardMarkup(True, True)
keyboardsend.row('Выгрузить резюме', 'Назад')

@bot.message_handler(['start'])
def start_message(message):
    if message.from_user.id == admin_id:
        bot.send_message(message.chat.id, f'Привет, {message.from_user.first_name}, Меню админа (Не забудь включить рассылку):', reply_markup=keyboardstartadmin)
    else:
        bot.send_message(message.chat.id, f'Привет, {message.from_user.first_name}',reply_markup=keyboardstart)

@bot.message_handler(content_types=['text'])
def send_message(message):
    global doc
    if message.text == 'Отправить резюме':
        bot.send_message(message.chat.id, 'Резюме отправлено на рассмотрение.', reply_markup=keyboardstart)
        doc = docx.Document()
        doc.add_heading('Резюме нового сотрудника:', 0)
        doc.add_paragraph(f'ФИО: {surname}, {name}, {lastname}')
        doc.add_paragraph(f'Телефон: {phone}')
        doc.add_paragraph(f'Дата рождения: {day}, {mounth}, {year}')
        doc.add_paragraph(f'Образование: {clas}')
        doc.add_paragraph(f'Последняя работа и причина увольнения: {lastwork}')
        doc.add_paragraph(f'Умения: {skill}')
        doc.add_paragraph(f'Почему хочет работать у нас: {whyy}')
        doc.add_paragraph(f'Желаемая зарплата: {sal}')
        doc.save('res.docx')
        send_document(message)
    if message.text == 'Мои страйки':
        bot.send_message(message.chat.id, 'Введите ФИО админа и дату смены в формате "Ф И О ДД ММ ГГ:')
        bot.register_next_step_handler(message, my)
    if message.text == 'Выгрузить страйки':
        send_strikes(message)
    if message.text == 'Страйки':
        bot.send_message(message.chat.id, 'Введите ФИО админа и дату смены в формате "Ф И О ДД ММ ГГ:')
        bot.register_next_step_handler(message, s_fio)
    if message.text == 'Назад':
        if message.from_user.id == admin_id:
            bot.send_message(message.chat.id, 'Главное меню:', reply_markup=keyboardstartadmin)
        else:
            bot.send_message(message.chat.id, 'Главное меню:', reply_markup=keyboardstart)
    if message.text == 'Резюме':
        bot.send_message(message.chat.id, 'Введите ФИО в формате "Фамилия Имя Отчество":')
        bot.register_next_step_handler(message, fio)
    if message.text == 'Выгрузить резюме':
        doc = docx.Document()
        doc.add_heading('Резюме нового сотрудника:', 0)
        doc.add_paragraph(f'ФИО: {surname}, {name}, {lastname}')
        doc.add_paragraph(f'Телефон: {phone}')
        doc.add_paragraph(f'Дата рождения: {day}, {mounth}, {year}')
        doc.add_paragraph(f'Образование: {clas}')
        doc.add_paragraph(f'Последняя работа и причина увольнения: {lastwork}')
        doc.add_paragraph(f'Умения: {skill}')
        doc.add_paragraph(f'Почему хочет работать у нас: {whyy}')
        doc.add_paragraph(f'Желаемая зарплата: {sal}')
        doc.save('res.docx')
        send_document(message)
    if message.text == 'Рассылка':
        bot.send_message(message.chat.id, 'Рассылки включены.', reply_markup=keyboardstartadmin)
        #Расписание начало дня
        schedule.every().day.at("10:47").do(check_smena)
        #Расписание обновление игр
        schedule.every().day.at("08:00").do(check_games)
        #Расписание "как смена?"
        schedule.every().day.at("03:00").do(check_worker)
        schedule.every().day.at("07:00").do(check_worker)
        schedule.every().day.at("15:00").do(check_worker)
        schedule.every().day.at("19:00").do(check_worker)
        schedule.every().day.at("22:00").do(check_worker)
        global mes
        mes = message
        while True:
            schedule.run_pending()
            time.sleep(1)


def check_smena():
    bot.send_message(mes.chat.id, f'Привет, {mes.from_user.first_name}\n'
                                  '1) Какие турниры в твоей смене? Напиши текстом.\n'
                                  '2) Сфоткай, как все красиво стоит в холодосах (2 фото)\n'
                                  '3) Прибери админку и отправь ее фото: чистый стол, нет лишних коробок и тп\n'
                                  '4) Скинь скрин графаны\n')
def check_worker():
    bot.send_message(mes.chat.id, f'Как смена, {mes.from_user.first_name}?\n'
                                  '1) Проверь чаты ВК\n'
                                  '2) Проверь холодильники, все должно стоять красиво\n'
                                  '3) Все свободные игровые места прибраны и без крошек\n'
                                  '4) Кинь фотки: общий зал, буткемп, мил зона (главное не слепи людей)\n'
                                  '5) Позвони с личного на рабочий +7(342)2-710-222, если вызов не прошел, тегни Тимофея')
def check_games():
    bot.send_message(mes.chat.id, f'{mes.from_user.first_name}, не спи!\n'
                                      'Проверь, обновлены ли следующие игры:\n'
                                      '- WOT\n'
                                      '- Warface\n'
                                      '- GTA V\n'
                                      '- LOL\n'
                                      '- CS:GO\n'
                                      '- Dota 2\n')


@bot.message_handler(content_types=['document'])
def send_strikes(message):
    docc = open(f'{f} {i} {o} {d} {m} {y}.docx', 'rb')
    bot.send_document(message.chat.id, docc)
def send_document(message):
    docc = open('res.docx', 'rb')
    bot.send_message(admin_id, f'Резюме от {message.from_user.first_name}')
    bot.send_document(admin_id, docc)
def my(message):
    ff,ii,oo, dd, mm, yy = message.text.split(" ")
    docc = open(f'{ff} {ii} {oo} {dd} {mm} {yy}.docx', 'rb')
    bot.send_document(message.chat.id, docc)
def s_fio(message):
    global strike
    global f, i, o, d, m, y
    doc = docx.Document()
    try:
        f, i, o, d, m, y = message.text.split(" ")
        doc.add_heading(f'Страйки сотрудника: {f} {i} {o}, за {d}.{m}.{y}', 0)
        bot.send_message(message.chat.id, 'Сотрудник обновил таблицу с играми?',reply_markup=keyboardanswer)
        bot.register_next_step_handler(message, one)
    except ValueError:
        bot.send_message(message.chat.id, 'Ошибка заполнения, введите по формату "Фамилия Имя Отчество День Месяц Год"')
        bot.message_handler(message, s_fio)
def one(message):
    if message.text == 'Да':
        strike = '2'
        doc.add_paragraph(f'Не обновил таблицу - {strike}')
        bot.send_message(message.chat.id, 'Ушел с админки и не поставил qr код', reply_markup=keyboardanswer)
        bot.register_next_step_handler(message, two)
    if message.text == 'Нет':
        strike = '0'
        doc.add_paragraph(f'Не обновил таблицу - {strike}')
        bot.send_message(message.chat.id, 'Ушел с админки и не поставил qr код', reply_markup=keyboardanswer)
        bot.register_next_step_handler(message, two)
def two(message):
    if message.text == 'Да':
        strike = '2'
        doc.add_paragraph(f'Ушел с админки и не поставил qr код - {strike}')
        bot.send_message(message.chat.id, 'скинул скрин Grafana с забитыми дисками / или не скинул Grafana', reply_markup=keyboardanswer)
        bot.register_next_step_handler(message, three)
    if message.text == 'Нет':
        strike = '0'
        doc.add_paragraph(f'Ушел с админки и не поставил qr код - {strike}')
        bot.send_message(message.chat.id, 'скинул скрин Grafana с забитыми дисками / или не скинул Grafana', reply_markup=keyboardanswer)
        bot.register_next_step_handler(message, three)
def three(message):
    if message.text == 'Да':
        strike = '2'
        doc.add_paragraph(f'скинул скрин Grafana с забитыми дисками / или не скинул Grafana - {strike}')
        bot.send_message(message.chat.id, 'не звонил на рабочий по требованию бота в тг', reply_markup=keyboardanswer)
        bot.register_next_step_handler(message, four)
    if message.text == 'Нет':
        strike = '0'
        doc.add_paragraph(f'скинул скрин Grafana с забитыми дисками / или не скинул Grafana - {strike}')
        bot.send_message(message.chat.id, 'не звонил на рабочий по требованию бота в тг ', reply_markup=keyboardanswer)
        bot.register_next_step_handler(message, four)
def four(message):
    if message.text == 'Да':
        strike = '2'
        doc.add_paragraph(f'не звонил на рабочий по требованию бота в тг - {strike}')
        bot.send_message(message.chat.id, 'скинул кривые фото зала', reply_markup=keyboardanswer)
        bot.register_next_step_handler(message, five)
    if message.text == 'Нет':
        strike = '0'
        doc.add_paragraph(f'не звонил на рабочий по требованию бота в тг - {strike}')
        bot.send_message(message.chat.id, "скинул кривые фото зала", reply_markup=keyboardanswer)
        bot.register_next_step_handler(message, five)
def five(message):
    if message.text == 'Да':
        strike = '2'
        doc.add_paragraph(f'скинул кривые фото зала - {strike}')
        bot.send_message(message.chat.id, 'скинул фото с куртками на креслах или не прибранным местом', reply_markup=keyboardanswer)
        bot.register_next_step_handler(message, six)
    if message.text == 'Нет':
        strike = '0'
        doc.add_paragraph(f'скинул кривые фото зала - {strike}')
        bot.send_message(message.chat.id, "скинул фото с куртками на креслах или не прибранным местом", reply_markup=keyboardanswer)
        bot.register_next_step_handler(message, six)
def six(message):
    if message.text == 'Да':
        strike = '2'
        doc.add_paragraph(f'скинул фото с куртками на креслах или не прибранным местом - {strike}')
        bot.send_message(message.chat.id, 'посторонние на админке', reply_markup=keyboardanswer)
        bot.register_next_step_handler(message, seven)
    if message.text == 'Нет':
        strike = '0'
        doc.add_paragraph(f'скинул фото с куртками на креслах или не прибранным местом - {strike}')
        bot.send_message(message.chat.id, "посторонние на админке", reply_markup=keyboardanswer)
        bot.register_next_step_handler(message, seven)
def seven(message):
    if message.text == 'Да':
        strike = '4'
        doc.add_paragraph(f'посторонние на админке - {strike}')
        bot.send_message(message.chat.id, 'на смене без мерча', reply_markup=keyboardanswer)
        bot.register_next_step_handler(message, e)
    if message.text == 'Нет':
        strike = '0'
        doc.add_paragraph(f'посторонние на админке - {strike}')
        bot.send_message(message.chat.id, "на смене без мерча", reply_markup=keyboardanswer)
        bot.register_next_step_handler(message, e)
def e(message):
    if message.text == 'Да':
        strike = '4'
        doc.add_paragraph(f'на смене без мерча - {strike}')
        bot.send_message(message.chat.id, 'оставил 1 полный бак с мусором', reply_markup=keyboardanswer)
        bot.register_next_step_handler(message, ee)
    if message.text == 'Нет':
        strike = '0'
        doc.add_paragraph(f'на смене без мерча - {strike}')
        bot.send_message(message.chat.id, "оставил 1 полный бак с мусором", reply_markup=keyboardanswer)
        bot.register_next_step_handler(message, ee)
def ee(message):
    if message.text == 'Да':
        strike = '4'
        doc.add_paragraph(f'оставил 1 полный бак с мусором - {strike}')
        bot.send_message(message.chat.id, 'не ответил на чат вк в течении 15 мин', reply_markup=keyboardanswer)
        bot.register_next_step_handler(message, eee)
    if message.text == 'Нет':
        strike = '0'
        doc.add_paragraph(f'оставил 1 полный бак с мусором - {strike}')
        bot.send_message(message.chat.id, "не ответил на чат вк в течении 15 мин", reply_markup=keyboardanswer)
        bot.register_next_step_handler(message, eee)
def eee(message):
    if message.text == 'Да':
        strike = '4'
        doc.add_paragraph(f'не ответил на чат вк в течении 15 мин - {strike}')
        bot.send_message(message.chat.id, 'задержка открытия смены более чем на 30 мин без уведомления в общем чате', reply_markup=keyboardanswer)
        bot.register_next_step_handler(message, eeee)
    if message.text == 'Нет':
        strike = '0'
        doc.add_paragraph(f'не ответил на чат вк в течении 15 мин - {strike}')
        bot.send_message(message.chat.id, "задержка открытия смены более чем на 30 мин без уведомления в общем чате", reply_markup=keyboardanswer)
        bot.register_next_step_handler(message, eeee)
def eeee(message):
    if message.text == 'Да':
        strike = '4'
        doc.add_paragraph(f'задержка открытия смены более чем на 30 мин без уведомления в общем чате - {strike}')
        bot.send_message(message.chat.id, 'в твоей смене обнаружена недостача по холодильнику', reply_markup=keyboardanswer)
        bot.register_next_step_handler(message, eeeee)
    if message.text == 'Нет':
        strike = '0'
        doc.add_paragraph(f'задержка открытия смены более чем на 30 мин без уведомления в общем чате - {strike}')
        bot.send_message(message.chat.id, "в твоей смене обнаружена недостача по холодильнику", reply_markup=keyboardanswer)
        bot.register_next_step_handler(message, eeeee)
def eeeee(message):
    if message.text == 'Да':
        strike = '5'
        doc.add_paragraph(f'в твоей смене обнаружена недостача по холодильнику - {strike}')
        bot.send_message(message.chat.id, 'не ответил на чат вк в течении 30 мин', reply_markup=keyboardanswer)
        bot.register_next_step_handler(message, eeeeee)
    if message.text == 'Нет':
        strike = '0'
        doc.add_paragraph(f'в твоей смене обнаружена недостача по холодильнику - {strike}')
        bot.send_message(message.chat.id, "не ответил на чат вк в течении 30 мин", reply_markup=keyboardanswer)
        bot.register_next_step_handler(message, eeeeee)
def eeeeee(message):
    if message.text == 'Да':
        strike = '8'
        doc.add_paragraph(f'не ответил на чат вк в течении 30 мин - {strike}')
        bot.send_message(message.chat.id, 'ночник соврал что обновил игры', reply_markup=keyboardanswer)
        bot.register_next_step_handler(message, eeeeeee)
    if message.text == 'Нет':
        strike = '0'
        doc.add_paragraph(f'не ответил на чат вк в течении 30 мин - {strike}')
        bot.send_message(message.chat.id, "ночник соврал что обновил игры", reply_markup=keyboardanswer)
        bot.register_next_step_handler(message, eeeeeee)
def eeeeeee(message):
    if message.text == 'Да':
        strike = '8'
        doc.add_paragraph(f'ночник соврал что обновил игры - {strike}')
        bot.send_message(message.chat.id, 'не продежурил', reply_markup=keyboardanswer)
        bot.register_next_step_handler(message, eeeeeeee)
    if message.text == 'Нет':
        strike = '0'
        doc.add_paragraph(f'ночник соврал что обновил игры - {strike}')
        bot.send_message(message.chat.id, "не продежурил", reply_markup=keyboardanswer)
        bot.register_next_step_handler(message, eeeeeeee)
def eeeeeeee(message):
    if message.text == 'Да':
        strike = '10'
        doc.add_paragraph(f'не продежурил - {strike}')
        bot.send_message(message.chat.id, 'не поставил #пс +номер или не пробил номер гостя на PS5', reply_markup=keyboardanswer)
        bot.register_next_step_handler(message, eeeeeeeee)
    if message.text == 'Нет':
        strike = '0'
        doc.add_paragraph(f'не продежурил - {strike}')
        bot.send_message(message.chat.id, "не поставил #пс +номер или не пробил номер гостя на PS5", reply_markup=keyboardanswer)
        bot.register_next_step_handler(message, eeeeeeeee)
def eeeeeeeee(message):
    if message.text == 'Да':
        strike = '10'
        doc.add_paragraph(f'не поставил #пс +номер или не пробил номер гостя на PS5 - {strike}')
        bot.send_message(message.chat.id, 'получил подтвержденный негативный отзыв клиента', reply_markup=keyboardanswer)
        bot.register_next_step_handler(message, eeeeeeeeee)
    if message.text == 'Нет':
        strike = '0'
        doc.add_paragraph(f'не поставил #пс +номер или не пробил номер гостя на PS5 - {strike}')
        bot.send_message(message.chat.id, "получил подтвержденный негативный отзыв клиента", reply_markup=keyboardanswer)
        bot.register_next_step_handler(message, eeeeeeeeee)
def eeeeeeeeee(message):
    if message.text == 'Да':
        strike = '10'
        doc.add_paragraph(f'получил подтвержденный негативный отзыв клиента - {strike}')
        bot.send_message(message.chat.id, 'уснул на смене', reply_markup=keyboardanswer)
        bot.register_next_step_handler(message, eeeeeeeeeee)
    if message.text == 'Нет':
        strike = '0'
        doc.add_paragraph(f'получил подтвержденный негативный отзыв клиента - {strike}')
        bot.send_message(message.chat.id, "уснул на смене", reply_markup=keyboardanswer)
        bot.register_next_step_handler(message, eeeeeeeeeee)
def eeeeeeeeeee(message):
    if message.text == 'Да':
        strike = '10'
        doc.add_paragraph(f'уснул на смене - {strike}')
        bot.send_message(message.chat.id, 'играл на смене на ПК, PS5', reply_markup=keyboardanswer)
        bot.register_next_step_handler(message, eeeeeeeeeeee)
    if message.text == 'Нет':
        strike = '0'
        doc.add_paragraph(f'уснул на смене - {strike}')
        bot.send_message(message.chat.id, "играл на смене на ПК, PS5", reply_markup=keyboardanswer)
        bot.register_next_step_handler(message, eeeeeeeeeeee)
def eeeeeeeeeeee(message):
    if message.text == 'Да':
        strike = '10'
        doc.add_paragraph(f'играл на смене на ПК, PS5 - {strike}')
        bot.send_message(message.chat.id, 'Страйки записаны', reply_markup=keyboardstrike)
        doc.save(f'{f} {i} {o} {d} {m} {y}.docx')
        bot.register_next_step_handler(message, send_strikes)
    if message.text == 'Нет':
        strike = '0'
        doc.add_paragraph(f'играл на смене на ПК, PS5 - {strike}')
        bot.send_message(message.chat.id, "Страйки записаны", reply_markup=keyboardstrike)
        doc.save(f'{f} {i} {o} {d} {m} {y}.docx')
        bot.register_next_step_handler(message, send_strikes)


def fio(message):
    global surname, name, lastname
    try:
        surname, name, lastname = message.text.split(" ")
        bot.send_message(message.chat.id, f'ФИО нового сотрудника: {surname} {name} {lastname}')
        bot.send_message(message.chat.id, 'Введите номер телефона в формате "89XXXXXXXXX"')
        bot.register_next_step_handler(message, number)
    except ValueError:
        bot.send_message(message.chat.id, 'Ошибка в заполнении, введите по формату')
        bot.register_next_step_handler(message, fio)
def number(message):
    global phone
    phone = message.text
    bot.send_message(message.chat.id, f'Телефон сотрудника: {phone}')
    bot.send_message(message.chat.id, 'Введите дату рождения (DD.MM.YY):')
    bot.register_next_step_handler(message, bd)
def bd(message):
    global day, mounth, year
    day, mounth, year = message.text.split(".")
    bot.send_message(message.chat.id, f'Дата рождения сотрудника: {day}.{mounth}.{year}')
    bot.send_message(message.chat.id, 'Введите Ваше образование:')
    bot.register_next_step_handler(message, classes)
def classes(message):
    global clas
    clas = message.text
    bot.send_message(message.chat.id, f'Ваше образование {clas}')
    bot.send_message(message.chat.id, f'Введите ваше прошлое место работы и причину увольнения:')
    bot.register_next_step_handler(message, last_work)
def last_work(message):
    global lastwork
    lastwork = message.text
    bot.send_message(message.chat.id, f'Ваше последнее место работы и причина увольнения: {lastwork}')
    bot.send_message(message.chat.id, 'Расскажи о себе, что умеешь?')
    bot.register_next_step_handler(message, skills)
def skills(message):
    global skill
    skill = message.text
    bot.send_message(message.chat.id, f'Ваше био: {skill}')
    bot.send_message(message.chat.id, 'Почему ты хочешь у нас работать?')
    bot.register_next_step_handler(message, why)
def why(message):
    global whyy
    whyy = message.text
    bot.send_message(message.chat.id, f'Причина, почему ты хочешь с нами работать -{whyy}')
    bot.send_message(message.chat.id, 'На какую зп рассчитываете?')
    bot.register_next_step_handler(message, sallary)
def sallary(message):
    global sal
    sal = message.text
    bot.send_message(message.chat.id, 'Анкета заполнена', reply_markup=keyboardboss)

print('all started successful')
bot.polling()